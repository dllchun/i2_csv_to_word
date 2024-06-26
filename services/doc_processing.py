from docx import Document
from docx.enum.section import WD_ORIENTATION
import pandas as pd
from datetime import datetime
from services.utils.file_utils import set_table_borders, add_bold_text
import tempfile


def generate_word_doc(combined_df: pd.DataFrame, language: str):
    doc = Document()

    # Set page orientation to landscape
    section = doc.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Add a table
    table = doc.add_table(rows=1, cols=len(combined_df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells

    column_names = [
        f"Ref#\n參考#",
        f"Smart Area\n智慧範疇",
        f"Initiative Group\n措施组合",
        f"Initiative\n措施",
        f"Progress Update\n進度",
        f"In progress / Completed/ Ongoing\n進行中/ 已完成 / 持續進行",
    ]
    for i, column in enumerate(column_names):
        hdr_cells[i].text = column

    # Add the data to the table
    for index, row in combined_df.iterrows():
        row_cells = table.add_row().cells
        for i, column in enumerate(combined_df.columns):
            cell_text = str(row[column])
            p = row_cells[i].paragraphs[0]
            if column == "Progress Update":
                add_bold_text(p, cell_text)
            else:
                row_cells[i].text = cell_text

    # Set table borders
    set_table_borders(table)

    # Save the document
    # doc_name = f"output/{datetime.now().strftime('%Y-%m-%d')}-{language}.docx"
    # doc.save(doc_name)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        return tmp.name

    print(f"Document saved as {doc_name}")
